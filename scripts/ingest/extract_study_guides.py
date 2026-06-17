"""Extract the instructor's per-unit study-guide questions.

Input:
  biol study guides/1592 Unit N Study Guide Questions.docx  (Units 1–17)

Each file is a flat list of open-ended study prompts written in the
instructor's own voice — these are NOT multiple-choice items to copy, but
*topic/coverage targets* the generated questions should mostly align with.

Output:
  content/study_guides.json — { "units": { "1": ["...", "..."], ... } }

The raw .docx live under "biol study guides/" (gitignored, like materials/);
this committed JSON is what build_units.py folds into each unit blob.

Idempotent. Re-run after the instructor revises a study guide.
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path

from docx import Document

ROOT = Path(__file__).resolve().parents[2]
SRC_DIR = ROOT / "biol study guides"
OUT = ROOT / "content" / "study_guides.json"

UNIT_RE = re.compile(r"Unit\s+(\d+)", re.IGNORECASE)


def extract_questions(doc: Document) -> list[str]:
    """Pull the study-prompt paragraphs, dropping the title line.

    The instructor's files put each prompt in a 'List Paragraph' style with
    the unit title as the first 'Normal' paragraph. We prefer List Paragraph
    items, but fall back to "every non-empty paragraph after the first" so a
    file with different styling still extracts cleanly.
    """
    list_items = [
        p.text.strip()
        for p in doc.paragraphs
        if p.text.strip() and (p.style.name or "").startswith("List")
    ]
    if list_items:
        return list_items
    # Fallback: skip the first non-empty paragraph (the title), keep the rest.
    nonempty = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return nonempty[1:] if len(nonempty) > 1 else nonempty


def main() -> int:
    if not SRC_DIR.exists():
        print(f"Missing source folder: {SRC_DIR}", file=sys.stderr)
        return 1

    docs = sorted(SRC_DIR.glob("*.docx"))
    if not docs:
        print(f"No .docx files in {SRC_DIR}", file=sys.stderr)
        return 1

    units: dict[str, list[str]] = {}
    for path in docs:
        m = UNIT_RE.search(path.stem)
        if not m:
            print(f"  SKIP  {path.name} — no 'Unit N' in filename", file=sys.stderr)
            continue
        unit_no = int(m.group(1))
        questions = extract_questions(Document(str(path)))
        if not questions:
            print(f"  WARN  unit {unit_no:02d}  no questions extracted from {path.name}")
            continue
        units[str(unit_no)] = questions
        print(f"  OK    unit {unit_no:02d}  {len(questions):>2} prompts  ({path.name})")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    with OUT.open("w", encoding="utf-8") as fp:
        json.dump({"units": units}, fp, ensure_ascii=False, indent=2)

    print(f"\nWrote {len(units)} unit(s) → {OUT.relative_to(ROOT)}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
